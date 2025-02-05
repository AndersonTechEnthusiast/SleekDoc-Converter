# importando tkinter 
import tkinter # Importa a biblioteca tkinter, que é usada para criar interfaces gráficas no Python
from tkinter import filedialog # Importa o módulo 'filedialog' de tkinter, que permite abrir um explorador de arquivos
from pathlib import Path # do 'pathlib' importe 'Path' 
import shutil # import ('ferramentas uteis para operações no sistema de arquivos (shell)')
from docx import Document # do docx importa Documento (com 'D' maisculo)
from docx.enum.text import WD_ALIGN_PARAGRAPH # do docx.enum.text (enumeradores de texto) , importa WD_ALIGN_PARAGRAPH (alinhadores do paragrafos do Word (WD))
from docx.enum.text import WD_COLOR_INDEX # do docx.enum.text (enumeradores de texto), importa WD_COLOR_INDEX  (cor de fundo do texto ``taxa de marcação``)
from docx.shared import RGBColor # RGBColor: Usado para definir cores no formato RGB (vermelho, verde, azul)
import time # time: Usado para manipular o tempo (como esperar um tempo específico)
from reportlab.lib import colors # colors: Contém cores predefinidas para usar no seu código (como 'red', 'blue')
from reportlab.lib.pagesizes import A4 # A4: Define o tamanho de página A4 para documentos em PDF
from reportlab.lib.colors import Color # Color: Usado para criar novas cores personalizadas no formato RGB ou CMYK
from reportlab.lib.styles import getSampleStyleSheet # getSampleStyleSheet: Obtém um conjunto de estilos padrão para usar em textos
from reportlab.lib.styles import ParagraphStyle # ParagraphStyle: Permite definir estilos personalizados para parágrafos
from reportlab.platypus import SimpleDocTemplate # SimpleDocTemplate: Usado para criar documentos PDF de forma simples
from reportlab.platypus import Paragraph # Paragraph: Usado para adicionar textos formatados no documento PDF
from reportlab.platypus import Image # Image: Usado para inserir imagens nos documentos PDF
from reportlab.lib.enums import TA_LEFT # TA_LEFT: Alinha o texto à esquerda no documento PDF
from reportlab.lib.enums import TA_CENTER # TA_CENTER: Alinha o texto ao centro no documento PDF
from reportlab.lib.enums import TA_RIGHT # TA_RIGHT: Alinha o texto à direita no documento PDF
from reportlab.lib.enums import TA_JUSTIFY # TA_JUSTIFY: Justifica o texto no documento PDF (distribui uniformemente)
from PIL import Image as Pillow # Pillow: Biblioteca para manipulação de imagens (como redimensionar ou converter formatos)
import pyautogui # pyautogui: Permite automação de ações no teclado e mouse (como clicar, digitar)
import keyboard # keyboard: Permite capturar e interagir com o teclado (como pressionar teclas)
import os # os: Permite interagir com o sistema operacional (como manipular arquivos e pastas)
import xml.etree.ElementTree # xml.etree.ElementTree: Usado para ler e manipular arquivos XML (como o DOCX)
# Cria a janela principal (root) para a interface gráfica, mas ainda não exibe
raiz = tkinter.Tk()
# Oculta a janela principal, já que não queremos mostrar nada, apenas o explorador de arquivos
raiz.withdraw()
# Abre o explorador de arquivos e permite ao usuário selecionar um arquivo
# O caminho do arquivo selecionado será armazenado na variável 'file_path'
file_path = filedialog.askopenfilename() # captura o arquivo selecionado
# converter Docx para PDF
while '.docx' not in file_path: # enquanto '.docx' não estiver contido no documento
    print("seu documento não é um Word (.docx)") # imprima a mensagem de error
    time.sleep(5) # espere 5 segundos
    raiz = tkinter.Tk() # refaça a janela Tkinter
    raiz.withdraw() # esconda a janela Tkinter
    file_path = filedialog.askopenfilename() # captura navamente o documento

Word_file = {} # cria uma array vazia para alocar todos os dados do documento
contador_from_paragraphs = 0 # contador para identificar os paragrafos
contador_for_comprimentos_paragrafos = -1

# cria uma nova instância da Classe requerida do 'docx'
Doc = Document(str(file_path)) # passa o nome do documento capturado como parametro

# acessa a sessão 'sections' do documento para pegar as margens
sections = Doc.sections[0]

margin_superior = sections.top_margin.pt # pega a margem superior do documento em pt (pontos)
margin_left = sections.left_margin.pt # pega a margem esquerda do documento em pt (pontos)
margin_right = sections.right_margin.pt # pega a margem direita do documento em pt (pontos)
margin_inferior = sections.bottom_margin.pt # pega a margem inferior do documento em pt (pontos)

# guarda tudo numa array 
ABNT_margins = {
    'superior': margin_superior,
    'direita': margin_right,
    'esquerda': margin_left,
    'inferior': margin_inferior
}

# cria uma array para armazenar seus comprimentos 
comprimentos_paragrafos = {}

# cria um índice incial
indice_comprimento_paragrafos = 0

# criando um contador
contador_xml = 0

# array para armazenar as imagens
imagens_rId = []

# cria array para armazenar dimensões cy e cx
dimensoes_cy_cx = {}

# contador 
contador_dimensoes_cy_cx = -1
# cria uma variavel (paragrafos)
# para cada paragrafo do documento (Doc.paragraphs) uma variavel paragrafos
for paragrafos in Doc.paragraphs:
    # print(paragrafos.style.name)
    # aciona o contador 
    contador_from_paragraphs += 1
    # captura o alinhamento de cada paragrafo
    alinhamento = paragrafos.alignment # alignment -> retorna o alinhamento do paragrafo em questão
    if alinhamento is None: 
        alinhamento = TA_LEFT
    if alinhamento == WD_ALIGN_PARAGRAPH.RIGHT: # WD_ALING_PARAGRAPH RETORNA A POSIÇÃO DO PARAGRAFO (.RIGHT,.CENTER,.LEFT,.JUSTIFY)
        alinhamento = TA_RIGHT
    if alinhamento == WD_ALIGN_PARAGRAPH.CENTER: # WD_ALING_PARAGRAPH RETORNA A POSIÇÃO DO PARAGRAFO (.RIGHT,.CENTER,.LEFT,.JUSTIFY)
        alinhamento = TA_CENTER
    if alinhamento == WD_ALIGN_PARAGRAPH.LEFT: # WD_ALING_PARAGRAPH RETORNA A POSIÇÃO DO PARAGRAFO (.RIGHT,.CENTER,.LEFT,.JUSTIFY)
        alinhamento = TA_LEFT
    if alinhamento == WD_ALIGN_PARAGRAPH.JUSTIFY: # WD_ALING_PARAGRAPH RETORNA A POSIÇÃO DO PARAGRAFO (.RIGHT,.CENTER,.LEFT,.JUSTIFY)
        alinhamento = TA_JUSTIFY
    
    # cria uma variavel para alocar o espaçamento
    padding_after = paragrafos.paragraph_format.space_after # variavel do paragrafo em questão 'paragrafos' , '.' acessa uma subcamanda do paragrafo ,'paragraph_format' formato do paragrafo , '.' acessa a subcamada do formato do paragrafo , 'space_after' espaçamento anterior
    padding_before = paragrafos.paragraph_format.space_before # variavel do paragrafo em questão 'paragrafos' , '.' acessa uma subcamanda do paragrafo ,'paragraph_format' formato do paragrafo , '.' acessa a subcamada do formato do paragrafo , 'space_after' espaçamento posterior
    if padding_after is None:
        padding_after = 0
    if padding_before is None:
        padding_before = 0
    if padding_after is not None:
        padding_after = (padding_after / 12700) + 10 + 5
    if padding_before is not None:
        padding_before = (padding_before / 12700) + 10 + 5
    # cria uma variavel para alocar o espaçamento entre linhas de cada linha
    espacamento_entre_linhas = paragrafos.paragraph_format.line_spacing # variavel do paragrafo em questão 'paragrafos' , '.' acessa uma subcamada do paragrafo , 'paragraph_format' , formato do paragrafo , '.' acessa a subcamada do formato do paragrafo , 'line_spacing' espaçamento de linha

    if espacamento_entre_linhas == None:
        espacamento_entre_linhas = 0

    # pegando o estilo da fonte 
    estilo = paragrafos.style.name.replace(' ','')
    
    contador_for_comprimentos_paragrafos = contador_for_comprimentos_paragrafos + 1

    # captura e aloca os comprimentos dos paragrafos
    comprimentos_paragrafos[contador_for_comprimentos_paragrafos] = len(paragrafos.text)

    # cria variavel para guardar o texto completo / comprimento 
    paragrafo_completo = ""

    # variavel para guardar a fonte do texto
    fonte = ""

    # HighLight_color (Marcação de Cor do Texto)
    highlight_color = ""
    
    # cria outra variavel 'styles'
    # para cada segmento ou trecho do paragrafo 'paragrafos (variavel criada)' uma variavel styles
    for styles in paragrafos.runs:
        # variavel fonte
        fonte           = styles.font.name # variavel criada 'styles' , '.' acesso o subcamada das runs 'styles' , 'font' fonte , '.' acessa uma subcamada da fonte , 'name' nome , busca o nome da fonte
        # variavel tamanho 
        tamanho         = styles.font.size # variavel criada 'styles' , '.' acesso a subcamada das runs 'styles' , 'font' fonte , '.' acesso uma subcamada da fonte , 'size' tamanho , 'busca o tamanho da fonte
        # variavel negrito
        negrito         = styles.bold # variavel criada 'styles , '.' acesso a subcamada das runs 'styles' , 'bold' negrito , verifica se o paragrafo tá em negrito (retorna apenas True ou None)
        # veriavel italico 
        italico         = styles.italic # variavel criada 'styles' , '.' acesso a subcamada das runs 'styles' , 'italic' italico verifica se o paragrafo é italico (returna True ou None)
        # variavel sublinado (underline)
        sublinhado      = styles.underline # variavel criada 'styles' , '.' acesso a subcamada das runs 'styles', 'underline' sublinha , verifica se o paragrafo contem underline (retorna True ou None)
        # variavel color 
        color           = styles.font.color.rgb # variavel criada 'styles', '.' acesso a subcamada das runs 'styles', 'font' acessa a subcamada fonte , '.' acessa a subcamada de 'font' , 'color' acessa a subcamada de cores , '.' acessa a subcamada de cores , 'rgb' puxa a cor em rgb
        # se a fonte for None
        if fonte is None:
            # sete 'Calibri(Body)'
            fonte = 'Calibri(Body)'
        if color is None: 
            color = RGBColor(0x00, 0x00, 0x00)
        if color is None and estilo != 'Normal':
            color = RGBColor(59, 89, 152)

        if tamanho is not None: 
            tamanho  = tamanho.pt
        
        # verifica se o texto tem hightlight_color , marcação de texto
        if styles.font.highlight_color: 
            highlight_color = styles.font.highlight_color
        
        #verifica as cores
        if highlight_color == WD_COLOR_INDEX.YELLOW:
            highlight_color = colors.yellow
        elif highlight_color == WD_COLOR_INDEX.GREEN:
            highlight_color = colors.green
        elif highlight_color == WD_COLOR_INDEX.WHITE:
            highlight_color = colors.white
        elif highlight_color == WD_COLOR_INDEX.BLUE:
            highlight_color = colors.blue
        elif highlight_color == WD_COLOR_INDEX.BLACK:
            highlight_color = colors.black
        elif highlight_color == WD_COLOR_INDEX.DARK_BLUE: 
            highlight_color = colors.darkblue
        elif highlight_color == WD_COLOR_INDEX.DARK_RED:
            highlight_color = colors.darkred
        elif highlight_color == WD_COLOR_INDEX.PINK:
            highlight_color = colors.pink
        elif highlight_color == WD_COLOR_INDEX.RED:
            highlight_color = colors.red
        elif highlight_color == WD_COLOR_INDEX.TEAL:
            highlight_color = colors.teal
        elif highlight_color == WD_COLOR_INDEX.TURQUOISE:
            highlight_color = colors.turquoise
        elif highlight_color == WD_COLOR_INDEX.VIOLET:
            highlight_color = colors.violet
        else:
            highlight_color = colors.white
    
        # variavel imagem 
        XML = None

        if "pic" in styles._r.xml:
            #
            # xml.etree.ElementTree (modulo importado do escopo)
            #
            # xml = Extensive Markup Language (Linguagem de Marcação Extensiva) , forma de organizar dados usando tags
            # 
            # etree = abreviação de Element Tree = Arvore de Elementos modulo do xml , estrutura em arvore
            # 
            # Element Tree = Arvore de Elementos (de novo) , entretanto é uma ferramenta trasmitida do modulo do etree (Arvore de Elementos) , ler , modificar e salvar
            # 
            # modulo é xml.etree.ElementTree = é um modulo importado do escopo do codigo
            # 
            # o segundo ElementTree de xml.etree.ElementTree.ElementTree() , .ElementTree() é a instância da classe ElementTree() , que fornece metodos de manipulação das ramificações da arvore , mas o que é uma arvore ? a melhor representação de arvore é a Liguagem de Marcação HMTL , como: 
            #
            # ARVORE:
            #
            # <html>
            #   <head>
            #       <title>
            #           titulo
            #        </title>
            #   </head>
            #   <body>
            #       <div>
            #           <p> texto </p>
            #       </div>
            #   </body>
            # </html>
            # isso acima é uma arvore: 
            # 
            # as Raizes são a tag <html>
            # as Ramificações são as tags <head> e <body> (elementos pais) são tags filhas da tag <html>
            # as Ramificações das Ramificações são <title> , <div> e <p> , são tags filhas da tag <head> , <title> tag filha da tag <head> , tags <div> e <p> filhas de <body>
            # e a instância da Classe .ElementTree() , trás métodos de como manipular essas tags
            # 
            # xml.etree.ElementTree () (modulo importado no escopo do codigo)
            # 
            # xml.etree.ElementTree.fromstring()
            # 
            # .fromstring() = significa 'do texto' , essa função converte o texto (string) XML (styles._r.xml) em uma arvore para que a instância da Classe .ElementTree() forneca os metodos necessarios para a arvore que era string mas foi tranformada em arvore
            #
            # variavel = xml.etree.ElementTree.ElementTree(xml.etree.ElementTree.fromstring(styles._r.xml)) -> Arvore (.fromstring()) com Métodos de Manipulação (.ElementTree())
            Arvore = xml.etree.ElementTree.ElementTree(xml.etree.ElementTree.fromstring(styles._r.xml))
            # variavel.getroot()
            #
            # getroot() , seria pegue a raiz da arvore
            #
            raiz = Arvore.getroot()
            # .iter() percorrer todos os elementos de acordo com o parametro
            #
            # .iter('parametro')
            # 
            # caso não souber qual parametro filtrar ou não saber como se escreve o parametro desejado faça o seguinte
            #
            # faça o seguinte escreva
            #
            # for blip in raiz.iter() // iter() sem parametros buscara todos os parametros da arvore
            #   print(blip) // mostrará tudo
            #
            # no meu caso eu queria a tag <a:blip r:embed...> do XML
            # 
            for blip in raiz.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                # após achar o blip que queria , após codar `for blip in raiz.iter()` iter() vazio , buscamos todos os caminhos completos: 
                # podemos aplicar blip.tag = .tag puxa a tag de todos os elementos das ramificações da raiz 
                # {http://schemas.openxmlformats.org/drawingml/2006/main}blip -> esse resultado é o blip que eu queria , ele se repetiu 5 vezes (tenho 5 imagens no meu documento) , todos são os mesmos
                # para achar o attributo que seria o embed 
                # <a:blip r:embed="rId">
                # blip = tag
                # embed = attribute
                # sendo assim: 
                # blip.tag = pega a tag
                # blip.attrib = pega o atributo da tag
                # print(f"tag: {blip.tag} | atributo: {blip.attrib}")
                # depois de achar você usar o
                # variavel (r_embed) = variavel (blip) , .attrib = atributo , .get = pegue , ou seja pegue o atributo
                r_embed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                # verifica se o atributo existe
                if r_embed:
                    # se existir , mostre-o
                    # print(f"embed encontrado: {r_embed}")
                    imagens_rId.append(r_embed)
                    XML = r_embed
            # procura por <wp:extent cx="5939790" cy="2671445"/> na arvore XML
            for extent in raiz.iter('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'):
                # incrementa um contador
                contador_dimensoes_cy_cx = contador_dimensoes_cy_cx + 1
                # adicionar o número do contador como chave da Array associativa
                dimensoes_cy_cx[contador_dimensoes_cy_cx] = extent.attrib
                # variavel (extent) é uma tupla [x,y] de [cy , cx]

        # variavel texto
        text = f'<p><span bgcolor="#{RGBColor(int(highlight_color.red * 255) , int(highlight_color.green * 255) , int(highlight_color.blue * 255))}"><font color = "#{color}">{styles.text}</font></span></p>' # variavel criada 'styles' , '.' acesso o subcamada da runs 'styles' , aloca o texto do paragrafo 'text' 

        # verifica se o texto é negrito
        if styles.bold:
            text = f'<b><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color = "#{color}">{styles.text}</font></span></b>'
        if styles.italic:
            text = f'<i><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color = "#{color}">{styles.text}</font></span></i>'
        if styles.underline:
            text = f'<u><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color = "#{color}">{styles.text}</font></span></u>'
        
        if styles.bold and styles.italic:
            text = f'<b><i><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color="#{color}">{styles.text}</font></span></i></b>'
        
        if styles.bold and styles.underline:
            text = f'<b><u><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color="#{color}">{styles.text}</font></span></u></b>'
        
        if styles.italic and styles.underline:
            text = f'<i><u><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color="#{color}">{styles.text}</font></span></u></i>'
        
        if styles.italic and styles.bold and styles.underline:
            text = f'<i><b><u><span bgcolor="#{RGBColor(int(highlight_color.red * 255),int(highlight_color.green * 255),int(highlight_color.blue * 255))}"><font color="#{color}">{styles.text}</font></span></u></b></i>'

        
        # aloca tudo na variavel 
        paragrafo_completo = paragrafo_completo + text
    indice_comprimento_paragrafos = indice_comprimento_paragrafos + 1

    # adiciona os valores na Array Associativa
    Word_file[f"parafrago-{contador_from_paragraphs}"] = {
        'alinhamento': alinhamento,
        'texto': paragrafo_completo,
        'fonte': fonte,
        'tamanho_fonte': tamanho,
        'negrito': negrito,
        'italico': italico,
        'sublinado': sublinhado,
        'fonte_cor': color,
        'espacamento_anterior': padding_before,
        'espacamento_posterior': padding_after,
        'espacamento_linhas': espacamento_entre_linhas,
        'estilo_do_paragrafo': estilo,
        'taxa_de_fundo_do_texto': highlight_color,
        'imagem': XML
    }

# cria a variavel (pasta) que alocará a Pasta 
# diretorio 
diretorio = "imagens" # aloca o nome da pasta na variavel 

#contador imagens 
contador_imagems = 0

# criando uma array para alocar todos os valores da imagem (guardar os caminhos da imagem)
caminhos_das_imagens = []
tamanhos_das_imagens = []
for rId in imagens_rId: # variavel (rId) para cada valor do imagens_rId
    # Doc.part.rels.values();
    # Doc = variavel que aloca a instância do documento Word (.docx)
    # part = referencia interna "pacote"
    # rels = é um dicionario que contém todos as "referencias" (links,imagens,videos,etc)
    # values() = pega todos os valores dentro do dicionario das referencias
    for target in Doc.part.rels.values():# cria uma variavel (target) para percorrer cada valor de referencia (links,videos,imagens,etc) , target para cada um dos valores de referencia
        # se existir "media" (mídia) , no target (variavel) , target_ref
        # target_ref - "caminho relativo"
        if "media" in target.target_ref:  # se media (mídia) existir , no caminho relativo
            # percorra toda a array dos rIds
            # se variavel (rId) que representa um dado da Array for igual ao target.rId da imagem
            if target.rId == rId: 
                # incrementa o contador 
                contador_imagems = contador_imagems + 1
                # transforma em jpg e manda para a pasta
                
                # pega a imagem em blob (binario)
                blob = target.target_part.blob

                # pega a extensão 
                extension = target.target_ref.split('.')[1]

                # nome da imagem 
                nome_da_imagem = f"{rId}-({contador_imagems}).{extension}"

                # caminho 
                caminho = os.path.join(diretorio , nome_da_imagem)
                
                # abre o documento e escreva o binario 
                with open(caminho , "wb") as arquivo:
                    arquivo.write(blob)
                
                # converte em jpg
                with Pillow.open(caminho) as arquivo:
                    arquivo = arquivo.convert("RGB")
                    width , height = arquivo.size
                    caminhos_das_imagens.append(f'{diretorio}/{rId}-{contador_imagems}.jpg')
                    tamanhos_das_imagens.append([width,height])
                    arquivo.save(f"{diretorio}/{rId}-{contador_imagems}.jpg" , "JPEG")

                
                # excluir duplicada
                if os.path.exists(caminho):
                    os.remove(caminho)

# array associativa para fazer as imagens concatenarem com seus caminhos
imagens_e_caminhos = {}
# contador de chaves 
contador_de_imagens_e_caminhos = -1
# um val para cada valor do array (caminhos_das_imagens)
for val in caminhos_das_imagens: # um val para cada 
    # incrementa o contador 
    contador_de_imagens_e_caminhos = contador_de_imagens_e_caminhos + 1
    # a chave da array associativa é o incrementador e é igual ao incrementador na array dos caminhos
    imagens_e_caminhos[contador_de_imagens_e_caminhos] = caminhos_das_imagens[contador_de_imagens_e_caminhos]

# nome do Documento PDF (baseado no documento Word)
documento_pdf = file_path.replace('.docx','.pdf').replace('D:/','')
# cria o documento PDF 
PDF = SimpleDocTemplate(f'./convertidos/{documento_pdf}',pagesize=A4)
# pega as estilizações PDF (styles)
styles = getSampleStyleSheet()
# cria uma lista para pegar todos os Paragrafos
paragrafos = []
# aplica as configurações de estilização do Paragrafo
contador_de_imagens_Trues = 0

for key , valor in Word_file.items():
    alinhamento = valor['alinhamento']
    texto = valor['texto']
    fonte = valor['fonte']
    tamanho_fonte = valor['tamanho_fonte']
    negrito = valor['negrito']
    italico = valor['italico']
    sublinhado = valor['sublinado']
    fonte_cor = valor['fonte_cor']
    espacamento_anterior = valor['espacamento_anterior']
    espacamento_posterior = valor['espacamento_posterior']
    espacamento_entre_linhas = valor['espacamento_linhas']
    estilo_do_paragrafo = valor['estilo_do_paragrafo']
    taxa_de_fundo_da_cor = valor['taxa_de_fundo_do_texto']
    imagem_do_paragrafo = valor['imagem']

    if espacamento_entre_linhas == 0:
        espacamento_entre_linhas = 1.5

    if taxa_de_fundo_da_cor == '':
        taxa_de_fundo_da_cor = None

    if imagem_do_paragrafo != None:
        
        # 914400 - > vem do EMU , English Metric Units , Unidades Metricas Englesa , diretamento do Microsoft Word Internacional para resoluções Metricas de Imagens e Dimensionamentos
        W_cx_cm = round((int(dimensoes_cy_cx[contador_de_imagens_Trues]['cx']) / 914400) * 2.54 , 2) # calculo Cm = Cx / EMU * 2.54 -> é o número de centimetros em polegadas
        H_cy_cm = round((int(dimensoes_cy_cx[contador_de_imagens_Trues]['cy']) / 914400) * 2.54 , 2) # calculo Cm = Cy / EMU * 2.54 -> é o número de centimetros em polegadas
        
        # regra de três para descobrir quanto vale 1 cm de px 
        # desfaz a tupla [x , y]
        largura,altura = A4
        # 21 cm === (largura) pixels
        # 1  cm === X pixels
        # X = (largura) * 1 / 21
        X_width = round((largura * 1) / 21 , 2)
        # 29.7 cm === (altura) pixels
        # 1    cm === X pixels
        # X = altura * 1 / 29.7
        X_height = round((altura * 1)/29.7 , 2)
        # print(f"tamanho em pixels: {tamanhos_das_imagens[contador_de_imagens_Trues][0]} pixels de largura \n tamanho em pixels: {tamanhos_das_imagens[contador_de_imagens_Trues][1]} pixels de altura \n tamanho em centimetros: {width_cm} cm de largura \n tamanho em centimetros: {height_cm} cm de altura \n tamanho em pixels (convertido): {X_width} pixels de largura \n tamanho em pixels (convertido): {X_height} pixels de altura")
        imagem = Image(caminhos_das_imagens[contador_de_imagens_Trues] , width=W_cx_cm * X_width , height=H_cy_cm * X_height)
        paragrafos.append(imagem)
        print(imagem_do_paragrafo)
        contador_de_imagens_Trues = contador_de_imagens_Trues + 1

    if estilo_do_paragrafo not in styles:
        estilo_do_paragrafo = 'Normal'

        # cria um estilo personalizado para cada paragrafoe
        style_paragraph = ParagraphStyle(
            # espaçamento anterior
            spaceBefore= int(espacamento_anterior),
            # nome do paragrafo 
            name=f"Documento",
            # fonte do texto do paragrafo 
            fontName = 'Helvetica',
            # tamanho da fonte do texto do paragrafo
            fontSize = styles[estilo_do_paragrafo].fontSize + 2,
            # espaçamento entre linhas (tamanho da fonte , vezes o espaçamento entre linhas)
            leading= styles[estilo_do_paragrafo].fontSize * espacamento_entre_linhas,
            # espaçamento posterior
            spaceAfter= int(espacamento_posterior),  
            # alinhamento   
            alignment=alinhamento,
            # Style do Paragrafo (Heading 1, Heading 2) 
            parent=styles[estilo_do_paragrafo]
        )
        # Adiciona os Paragrafos estilizados a Lista 
        paragrafos.append(Paragraph(texto, style_paragraph))
    else:
        # cria um estilo personalizado para cada paragrafoe
        style_paragraph = ParagraphStyle(
            # espaçamento anterior
            spaceBefore= int(espacamento_anterior),
            # nome do paragrafo 
            name=f"Documento",
            # fonte do texto do paragrafo 
            fontName = 'Helvetica',
            # tamanho da fonte do texto do paragrafo
            fontSize = tamanho_fonte,
            # espaçamento entre linhas (tamanho da fonte , vezes o espaçamento entre linhas)
            leading= tamanho_fonte * espacamento_entre_linhas,
            # espaçamento posterior
            spaceAfter= int(espacamento_posterior),  
            # alinhamento   
            alignment=alinhamento,
            # Style do Paragrafo (Heading 1, Heading 2) 
            parent=styles[estilo_do_paragrafo]
        )
        # Adiciona os Paragrafos estilizados a Lista 
        paragrafos.append(Paragraph(texto, style_paragraph))
    
try:
    # gera o Arquivo PDF
    PDF.build(paragrafos)
except Exception as e:
    print(f"Error : {e}")